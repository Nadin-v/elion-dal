"""Unit-тест парсинга документов (DOCX). PDF не генерируем — покрыт сидингом «Базы знаний»."""

from __future__ import annotations

from pathlib import Path

import docx

from elion_dal.ingestion.loaders import load_document, load_docx


def _make_docx(path: Path) -> None:
    d = docx.Document()
    d.add_paragraph("Положение об олимпиаде")
    d.add_paragraph("Пункт 1. Общие положения.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Дата"
    table.rows[0].cells[1].text = "15 марта"
    d.save(str(path))


def test_load_docx_paragraphs_and_tables(tmp_path):
    p = tmp_path / "doc.docx"
    _make_docx(p)
    text = load_docx(p)
    assert "Положение об олимпиаде" in text
    assert "Пункт 1" in text
    assert "Дата" in text and "15 марта" in text  # ячейки таблицы тоже извлекаются


def test_load_document_dispatch_by_extension(tmp_path):
    p = tmp_path / "doc.docx"
    _make_docx(p)
    assert "Положение" in load_document(p)

    import pytest

    with pytest.raises(ValueError):
        load_document(tmp_path / "file.rtf")
